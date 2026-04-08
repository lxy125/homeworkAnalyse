import json
import os
import re
import shutil
import subprocess
import tempfile
import time
from datetime import datetime
from pathlib import Path
from typing import Any

import httpx
import streamlit as st
from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import XmlPart
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from dotenv import load_dotenv
from openai import OpenAI
from openpyxl import load_workbook
from pypdf import PdfReader

load_dotenv()

DEFAULT_OPENAI_BASE_URL = "https://ark.cn-beijing.volces.com/api/coding/v3"
DEFAULT_ANTHROPIC_BASE_URL = "https://ark.cn-beijing.volces.com/api/coding"
DEFAULT_ARK_FILE_BASE_URL = "https://ark.cn-beijing.volces.com/api/v3"
DEFAULT_MODEL = "doubao-seed-2.0-pro"
QUESTION_EXT = {".doc", ".docx", ".pdf", ".xls", ".xlsx"}
STUDENT_WORD_EXT = {".doc", ".docx"}
FILE_API_SUPPORTED_EXT = {
    ".pdf",
    ".docx",
    ".jpg",
    ".jpeg",
    ".png",
    ".gif",
    ".webp",
    ".bmp",
    ".tiff",
    ".ico",
    ".icns",
    ".sgi",
    ".jp2",
    ".heic",
    ".heif",
    ".mp4",
    ".avi",
    ".mov",
}
OUTPUT_DIR = Path("outputs")
OUTPUT_DIR.mkdir(exist_ok=True)
CLASS_ID = "class-demo"
COURSE_ID = "course_oop_design"
SOURCE_TYPE = "assignment"
EVENT_TYPE = "GRADING_RESULT_RECEIVED"
DEFAULT_SOURCE_ID = "assignment-001"
DEFAULT_MEMBER4_INGEST_URL = "http://127.0.0.1:8007/api/v1/analytics/events/ingest"
COURSE_KNOWLEDGE_POINTS = [
    {"id": "kp_object_oriented_basics", "name": "面向对象基础"},
    {"id": "kp_class_and_object", "name": "类与对象"},
    {"id": "kp_encapsulation_and_access_control", "name": "封装与访问控制"},
    {"id": "kp_inheritance", "name": "继承"},
    {"id": "kp_polymorphism", "name": "多态"},
    {"id": "kp_abstract_class_and_interface", "name": "抽象类与接口"},
    {"id": "kp_exception_handling", "name": "异常处理"},
    {"id": "kp_common_collections_and_generics", "name": "常用集合与泛型"},
]


def now_suffix() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def normalize_student_id(student_id: str) -> str:
    value = student_id.strip()
    if not value:
        raise ValueError("学生ID不能为空。")
    normalized = re.sub(r"[^0-9A-Za-z_-]+", "_", value).strip("_")
    if not normalized:
        raise ValueError("学生ID仅支持字母、数字、下划线和中划线。")
    return normalized[:64]


def normalize_text(text: str) -> str:
    cleaned = text.replace("\r", "\n")
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def extract_json(text: str) -> dict[str, Any]:
    text = text.strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    match = re.search(r"\{[\s\S]*\}", text)
    if not match:
        return {}

    try:
        return json.loads(match.group(0))
    except json.JSONDecodeError:
        return {}


def call_openai_compatible(api_key: str, base_url: str, model: str, system_prompt: str, user_prompt: str) -> str:
    client = OpenAI(api_key=api_key, base_url=base_url)
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.2,
        max_tokens=1800,
    )
    return (response.choices[0].message.content or "").strip()


def call_anthropic_compatible(api_key: str, base_url: str, model: str, system_prompt: str, user_prompt: str) -> str:
    url = base_url.rstrip("/")
    if not url.endswith("/v1/messages"):
        url = f"{url}/v1/messages"

    headers = {
        "x-api-key": api_key,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json",
    }
    payload = {
        "model": model,
        "max_tokens": 1800,
        "temperature": 0.2,
        "system": system_prompt,
        "messages": [{"role": "user", "content": [{"type": "text", "text": user_prompt}]}],
    }

    with httpx.Client(timeout=120.0) as client:
        response = client.post(url, headers=headers, json=payload)
        response.raise_for_status()
        data = response.json()

    parts = data.get("content", [])
    texts = [part.get("text", "") for part in parts if part.get("type") == "text"]
    return "\n".join([t.strip() for t in texts if t.strip()]).strip()


def call_model(protocol: str, api_key: str, base_url: str, model: str, system_prompt: str, user_prompt: str) -> str:
    if protocol == "Anthropic兼容":
        return call_anthropic_compatible(api_key, base_url, model, system_prompt, user_prompt)
    return call_openai_compatible(api_key, base_url, model, system_prompt, user_prompt)


def _extract_response_text(response: Any) -> str:
    text = (getattr(response, "output_text", "") or "").strip()
    if text:
        return text

    chunks: list[str] = []
    for output_item in getattr(response, "output", []) or []:
        for content in getattr(output_item, "content", []) or []:
            content_type = getattr(content, "type", "")
            if content_type in {"output_text", "text"}:
                value = getattr(content, "text", "") or ""
                if value.strip():
                    chunks.append(value.strip())
    return "\n".join(chunks).strip()


def resolve_file_api_base_url(base_url: str) -> str:
    env_url = os.getenv("ARK_FILE_BASE_URL", "").strip()
    if env_url:
        return env_url

    normalized = base_url.rstrip("/")
    if normalized.endswith("/api/coding/v3"):
        return normalized[: -len("/api/coding/v3")] + "/api/v3"
    if normalized.endswith("/api/coding"):
        return normalized[: -len("/api/coding")] + "/api/v3"
    return DEFAULT_ARK_FILE_BASE_URL


def _file_status(meta: Any) -> str:
    status = getattr(meta, "status", None)
    if isinstance(status, str) and status:
        return status
    if isinstance(meta, dict):
        raw = meta.get("status")
        if isinstance(raw, str):
            return raw
    return ""


def split_file_api_supported(paths: list[Path]) -> tuple[list[Path], list[Path]]:
    supported: list[Path] = []
    unsupported: list[Path] = []
    for p in paths:
        if p.suffix.lower() in FILE_API_SUPPORTED_EXT:
            supported.append(p)
        else:
            unsupported.append(p)
    return supported, unsupported


def call_openai_compatible_with_files(
    api_key: str,
    base_url: str,
    model: str,
    system_prompt: str,
    user_prompt: str,
    file_paths: list[Path],
) -> str:
    file_base_url = resolve_file_api_base_url(base_url)
    client = OpenAI(api_key=api_key, base_url=file_base_url)
    uploaded_file_ids: list[str] = []
    try:
        for file_path in file_paths:
            with file_path.open("rb") as f:
                uploaded = client.files.create(file=(file_path.name, f), purpose="user_data")
                uploaded_file_ids.append(uploaded.id)
            deadline = time.time() + 300
            while True:
                meta = client.files.retrieve(uploaded.id)
                status = _file_status(meta)
                if status == "active":
                    break
                if status in {"failed", "error"}:
                    raise RuntimeError(f"文件预处理失败: {file_path.name}")
                if time.time() >= deadline:
                    raise TimeoutError(f"等待文件预处理超时: {file_path.name}")
                time.sleep(1.0)

        user_content = [{"type": "input_text", "text": user_prompt}]
        for fid in uploaded_file_ids:
            user_content.append({"type": "input_file", "file_id": fid})

        response = client.responses.create(
            model=model,
            temperature=0.2,
            max_output_tokens=1800,
            input=[
                {"role": "system", "content": [{"type": "input_text", "text": system_prompt}]},
                {"role": "user", "content": user_content},
            ],
        )
        return _extract_response_text(response)
    finally:
        for fid in uploaded_file_ids:
            try:
                client.files.delete(fid)
            except Exception:
                continue


def is_file_input_unsupported_error(exc: Exception) -> bool:
    text = str(exc).lower()
    keywords = ["404", "/files", "notfound", "unsupported", "input_file", "file type not supported"]
    return any(k in text for k in keywords)


def convert_legacy_office_file(file_path: Path, target_ext: str) -> Path:
    def _powershell_escape(path: str) -> str:
        return path.replace("'", "''")

    def _convert_with_windows_office(src: Path, dst: Path, dst_ext: str) -> bool:
        # Windows 下使用已安装的 Office 进行转换，作为 soffice 不可用时的兜底。
        src_esc = _powershell_escape(str(src.resolve()))
        dst_esc = _powershell_escape(str(dst.resolve()))
        if dst_ext == ".docx":
            script = f"""
$ErrorActionPreference = 'Stop'
$word = $null
$doc = $null
try {{
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $doc = $word.Documents.Open('{src_esc}')
  $doc.SaveAs2('{dst_esc}', 16)
}} finally {{
  if ($doc -ne $null) {{ $doc.Close([ref]$false) }}
  if ($word -ne $null) {{ $word.Quit() }}
}}
""".strip()
        elif dst_ext == ".xlsx":
            script = f"""
$ErrorActionPreference = 'Stop'
$excel = $null
$wb = $null
try {{
  $excel = New-Object -ComObject Excel.Application
  $excel.Visible = $false
  $excel.DisplayAlerts = $false
  $wb = $excel.Workbooks.Open('{src_esc}')
  $wb.SaveAs('{dst_esc}', 51)
}} finally {{
  if ($wb -ne $null) {{ $wb.Close($false) }}
  if ($excel -ne $null) {{ $excel.Quit() }}
}}
""".strip()
        else:
            return False

        result = subprocess.run(
            ["powershell", "-NoProfile", "-NonInteractive", "-Command", script],
            capture_output=True,
            text=True,
            check=False,
        )
        return result.returncode == 0 and dst.exists()

    soffice = shutil.which("soffice")
    temp_dir = Path(tempfile.mkdtemp(prefix="office_convert_"))
    convert_to = target_ext.lstrip(".")
    converted = temp_dir / f"{file_path.stem}.{convert_to}"
    if soffice:
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", convert_to, "--outdir", str(temp_dir), str(file_path)],
            capture_output=True,
            text=True,
            check=False,
        )
        if result.returncode == 0 and converted.exists():
            return converted

    if os.name == "nt" and _convert_with_windows_office(file_path, converted, target_ext):
        return converted

    raise ValueError(
        f"文件 `{file_path.name}` 为旧格式 `{file_path.suffix}`，当前环境无法直接解析。"
        f"请改传 `{target_ext}`，或安装 LibreOffice（soffice）/本机 Office 组件后自动转换。"
    )


def ensure_docx(file_path: Path) -> Path:
    ext = file_path.suffix.lower()
    if ext == ".docx":
        return file_path
    if ext == ".doc":
        return convert_legacy_office_file(file_path, ".docx")
    raise ValueError(f"不支持的Word文件类型: {ext}")


def ensure_xlsx(file_path: Path) -> Path:
    ext = file_path.suffix.lower()
    if ext == ".xlsx":
        return file_path
    if ext == ".xls":
        return convert_legacy_office_file(file_path, ".xlsx")
    raise ValueError(f"不支持的Excel文件类型: {ext}")


def extract_text_from_pdf(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    text = []
    for page in reader.pages:
        text.append(page.extract_text() or "")
    return normalize_text("\n".join(text))


def extract_text_from_word(file_path: Path) -> str:
    docx_path = ensure_docx(file_path)
    doc = Document(str(docx_path))
    lines = []
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            lines.append(txt)
    return normalize_text("\n".join(lines))


def extract_text_from_excel(file_path: Path) -> str:
    xlsx_path = ensure_xlsx(file_path)
    wb = load_workbook(str(xlsx_path), data_only=True, read_only=True)
    try:
        lines = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                row_values = []
                for val in row:
                    if val is not None and str(val).strip():
                        row_values.append(str(val).strip())
                if row_values:
                    lines.append(f"[{ws.title}] " + " | ".join(row_values))
        return normalize_text("\n".join(lines))
    finally:
        wb.close()


def extract_text(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext in {".doc", ".docx"}:
        return extract_text_from_word(file_path)
    if ext in {".xls", ".xlsx"}:
        return extract_text_from_excel(file_path)
    raise ValueError(f"不支持的文件类型: {ext}")


def generate_annotation_plan(
    protocol: str,
    api_key: str,
    base_url: str,
    model: str,
    question_text: str,
    student_segments: list[dict[str, Any]],
    segment_hint: str,
    reference_text: str = "",
    question_file_paths: list[Path] | None = None,
    reference_file_path: Path | None = None,
    student_file_path: Path | None = None,
    use_model_file_inputs: bool = False,
) -> dict[str, Any]:
    system_prompt = (
        "你是严谨且鼓励式的老师。你只返回JSON，不要返回Markdown或解释。"
        "批注要简短，聚焦知识点与改进建议。"
    )

    serialized = []
    for item in student_segments:
        serialized.append(f"index={item['index']} | content={item['text']}")

    if use_model_file_inputs:
        material_names = [p.name for p in (question_file_paths or [])]
        if reference_file_path:
            material_names.append(reference_file_path.name)
        if student_file_path and student_file_path.suffix.lower() in FILE_API_SUPPORTED_EXT:
            material_names.append(student_file_path.name)
        materials_desc = "、".join(material_names) if material_names else "未附加"
        user_prompt = f"""
请依据题目与学生作答给出批注计划。

输出JSON格式必须是：
{{
  "overall": "总评，不超过80字",
  "items": [
    {{"index": 1, "comment": "该段批注，不超过40字"}}
  ]
}}

要求：
1) items 只包含需要批注的条目，最多20条。
2) index 必须来自我提供的 index。
3) comment 使用中文。
4) 不要输出任何JSON以外文本。
5) 题目、材料、参考样例和学生作答文档已经作为文件附件提供：{materials_desc}。

学生作答条目（{segment_hint}）：
{chr(10).join(serialized)}
""".strip()
    else:
        user_prompt = f"""
请依据题目与学生作答给出批注计划。

输出JSON格式必须是：
{{
  "overall": "总评，不超过80字",
  "items": [
    {{"index": 1, "comment": "该段批注，不超过40字"}}
  ]
}}

要求：
1) items 只包含需要批注的条目，最多20条。
2) index 必须来自我提供的 index。
3) comment 使用中文。
4) 不要输出任何JSON以外文本。

题目：
{question_text}

参考老师批改样例（可选）：
{reference_text if reference_text else '无'}

学生作答条目（{segment_hint}）：
{chr(10).join(serialized)}
""".strip()
    if use_model_file_inputs and protocol != "Anthropic兼容":
        attach_paths = list(question_file_paths or [])
        if reference_file_path:
            attach_paths.append(reference_file_path)
        if student_file_path and student_file_path.suffix.lower() in FILE_API_SUPPORTED_EXT:
            attach_paths.append(student_file_path)
        if attach_paths:
            raw = call_openai_compatible_with_files(api_key, base_url, model, system_prompt, user_prompt, attach_paths)
        else:
            raw = call_model(protocol, api_key, base_url, model, system_prompt, user_prompt)
    else:
        raw = call_model(protocol, api_key, base_url, model, system_prompt, user_prompt)
    parsed = extract_json(raw)
    if not parsed:
        return {"overall": "整体完成较认真，建议根据题目要求进一步完善关键步骤。", "items": []}

    overall = str(parsed.get("overall", "")).strip() or "整体完成较认真，建议根据题目要求进一步完善关键步骤。"
    items = parsed.get("items", [])
    normalized_items = []
    for item in items:
        try:
            idx = int(item.get("index"))
            comment = str(item.get("comment", "")).strip()
            if idx > 0 and comment:
                normalized_items.append({"index": idx, "comment": comment[:120]})
        except Exception:
            continue

    return {"overall": overall[:120], "items": normalized_items}


def _coerce_bool(value: Any) -> bool | None:
    if isinstance(value, bool):
        return value
    if isinstance(value, int):
        return bool(value)
    if isinstance(value, str):
        raw = value.strip().lower()
        if raw in {"true", "1", "yes", "y"}:
            return True
        if raw in {"false", "0", "no", "n"}:
            return False
    return None


def _normalize_knowledge_point_results(raw_results: Any) -> list[dict[str, Any]]:
    allowed_ids = {item["id"] for item in COURSE_KNOWLEDGE_POINTS}
    if not isinstance(raw_results, list):
        return []

    normalized: list[dict[str, Any]] = []
    seen_ids: set[str] = set()
    for item in raw_results:
        if not isinstance(item, dict):
            continue
        kp_id = str(item.get("knowledge_point_id", "")).strip()
        is_correct = _coerce_bool(item.get("is_correct"))
        if kp_id not in allowed_ids or is_correct is None or kp_id in seen_ids:
            continue
        seen_ids.add(kp_id)
        normalized.append({"knowledge_point_id": kp_id, "is_correct": is_correct})
    return normalized


def generate_knowledge_point_results(
    protocol: str,
    api_key: str,
    base_url: str,
    model: str,
    student_id: str,
    question_text: str,
    student_text: str,
    reference_text: str = "",
    question_file_paths: list[Path] | None = None,
    reference_file_path: Path | None = None,
    use_model_file_inputs: bool = False,
) -> dict[str, Any]:
    system_prompt = (
        "你是教学评估专家。你只返回JSON，不要返回Markdown或解释。"
        "请严格根据题目和作答判断涉及知识点及是否基本正确。"
    )
    course_schema = {
        "course_id": COURSE_ID,
        "course_name": "面向对象程序设计",
        "knowledge_points": COURSE_KNOWLEDGE_POINTS,
    }
    if use_model_file_inputs:
        material_names = [p.name for p in (question_file_paths or [])]
        if reference_file_path:
            material_names.append(reference_file_path.name)
        materials_desc = "、".join(material_names) if material_names else "未附加"
        user_prompt = f"""
请根据题目内容、学生作答内容、课程知识点列表，判断学生作答涉及了哪些知识点，并判断每个知识点是否基本回答正确。

输出JSON格式必须是：
{{
  "results": [
    {{
      "knowledge_point_id": "必须是下方知识点id之一",
      "is_correct": true
    }}
  ]
}}

要求：
1) knowledge_point_id 只能从下面列表中选择，禁止输出列表之外的id。
2) is_correct 必须是布尔值 true 或 false。
3) 仅输出学生作答中涉及到的知识点。
4) 不要输出任何JSON以外文本。
5) 题目、材料和参考样例已经作为文件附件提供：{materials_desc}。

课程知识点列表：
{json.dumps(course_schema, ensure_ascii=False, indent=2)}

学生作答：
{student_text[:8000]}
""".strip()
    else:
        user_prompt = f"""
请根据题目内容、学生作答内容、课程知识点列表，判断学生作答涉及了哪些知识点，并判断每个知识点是否基本回答正确。

输出JSON格式必须是：
{{
  "results": [
    {{
      "knowledge_point_id": "必须是下方知识点id之一",
      "is_correct": true
    }}
  ]
}}

要求：
1) knowledge_point_id 只能从下面列表中选择，禁止输出列表之外的id。
2) is_correct 必须是布尔值 true 或 false。
3) 仅输出学生作答中涉及到的知识点。
4) 不要输出任何JSON以外文本。

题目与教师材料：
{question_text[:8000]}

参考样例（可选）：
{reference_text[:3000] if reference_text else '无'}

课程知识点列表：
{json.dumps(course_schema, ensure_ascii=False, indent=2)}

学生作答：
{student_text[:8000]}
""".strip()
    if use_model_file_inputs and protocol != "Anthropic兼容":
        attach_paths = list(question_file_paths or [])
        if reference_file_path:
            attach_paths.append(reference_file_path)
        if attach_paths:
            raw = call_openai_compatible_with_files(api_key, base_url, model, system_prompt, user_prompt, attach_paths)
        else:
            raw = call_model(protocol, api_key, base_url, model, system_prompt, user_prompt)
    else:
        raw = call_model(protocol, api_key, base_url, model, system_prompt, user_prompt)
    parsed = extract_json(raw)
    results = _normalize_knowledge_point_results(parsed.get("results", []) if parsed else [])
    return {
        "user_id": student_id,
        "class_id": CLASS_ID,
        "course_id": COURSE_ID,
        "event_type": EVENT_TYPE,
        "payload": {
            "source_id": DEFAULT_SOURCE_ID,
            "source_type": SOURCE_TYPE,
            "results": results,
        },
    }


def post_member4_event(
    ingest_url: str,
    user_id: str,
    source_id: str,
    results: list[dict[str, Any]],
) -> dict[str, Any]:
    event = {
        "user_id": user_id,
        "class_id": CLASS_ID,
        "course_id": COURSE_ID,
        "event_type": EVENT_TYPE,
        "payload": {
            "source_id": source_id,
            "source_type": SOURCE_TYPE,
            "results": results,
        },
    }
    request_body = {"events": [event]}
    with httpx.Client(timeout=20.0) as client:
        response = client.post(ingest_url, json=request_body)
        response.raise_for_status()
        content_type = response.headers.get("content-type", "")
        response_body: Any = response.text
        if "application/json" in content_type.lower():
            try:
                response_body = response.json()
            except Exception:
                response_body = response.text
        return {"request": request_body, "status_code": response.status_code, "response_body": response_body}


def get_or_add_comments_part(document: Document) -> XmlPart:
    doc_part = document.part
    for rel in doc_part.rels.values():
        if rel.reltype == RT.COMMENTS:
            return rel.target_part

    comments_xml = parse_xml(f"<w:comments {nsdecls('w')}/>")
    comments_part = XmlPart(PackURI("/word/comments.xml"), CT.WML_COMMENTS, comments_xml, doc_part.package)
    doc_part.relate_to(comments_part, RT.COMMENTS)
    return comments_part


def next_comment_id(comments_part: XmlPart) -> int:
    comments_root = comments_part.element
    max_id = -1
    for node in comments_root.findall(qn("w:comment")):
        raw = node.get(qn("w:id"))
        if raw is None:
            continue
        try:
            max_id = max(max_id, int(raw))
        except ValueError:
            continue
    return max_id + 1


def add_comment_to_paragraph(document: Document, paragraph: Any, comment_text: str, author: str = "AI批改助手") -> None:
    text = comment_text.strip()
    if not text:
        return

    comments_part = get_or_add_comments_part(document)
    cid = next_comment_id(comments_part)

    comment = OxmlElement("w:comment")
    comment.set(qn("w:id"), str(cid))
    comment.set(qn("w:author"), author)
    comment.set(qn("w:initials"), "AI")
    comment.set(qn("w:date"), datetime.utcnow().replace(microsecond=0).isoformat() + "Z")

    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    comment.append(p)
    comments_part.element.append(comment)

    p_elm = paragraph._p
    runs = [child for child in p_elm.iterchildren() if child.tag == qn("w:r")]
    if not runs:
        paragraph.add_run(" ")
        runs = [child for child in p_elm.iterchildren() if child.tag == qn("w:r")]

    if not runs:
        return

    first_run = runs[0]
    last_run = runs[-1]

    range_start = OxmlElement("w:commentRangeStart")
    range_start.set(qn("w:id"), str(cid))
    first_run.addprevious(range_start)

    range_end = OxmlElement("w:commentRangeEnd")
    range_end.set(qn("w:id"), str(cid))
    last_run.addnext(range_end)

    ref_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "CommentReference")
    rpr.append(rstyle)
    ref_run.append(rpr)
    cref = OxmlElement("w:commentReference")
    cref.set(qn("w:id"), str(cid))
    ref_run.append(cref)
    range_end.addnext(ref_run)


def annotate_word(
    student_path: Path,
    student_id: str,
    output_dir: Path,
    protocol: str,
    api_key: str,
    base_url: str,
    model: str,
    question_text: str,
    reference_text: str,
    question_file_paths: list[Path] | None = None,
    reference_file_path: Path | None = None,
    use_model_file_inputs: bool = False,
) -> tuple[Path, str, str]:
    docx_path = ensure_docx(student_path)
    doc = Document(str(docx_path))

    segments = []
    segment_map: dict[int, tuple[Any, str]] = {}
    idx = 1
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            segments.append({"index": idx, "text": txt})
            segment_map[idx] = (para, txt)
            idx += 1

    plan = generate_annotation_plan(
        protocol,
        api_key,
        base_url,
        model,
        question_text,
        segments,
        "Word段落",
        reference_text,
        question_file_paths=question_file_paths,
        reference_file_path=reference_file_path,
        student_file_path=docx_path,
        use_model_file_inputs=use_model_file_inputs,
    )

    if doc.paragraphs:
        add_comment_to_paragraph(doc, doc.paragraphs[0], f"总评：{plan['overall']}")

    for item in plan["items"]:
        mapped = segment_map.get(item["index"])
        if not mapped:
            continue
        para, _ = mapped
        add_comment_to_paragraph(doc, para, item["comment"])

    output_path = output_dir / f"{student_path.stem}-学生ID-{student_id}-批改后-{now_suffix()}.docx"
    output_dir.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path.resolve()))
    student_text = normalize_text("\n".join([item["text"] for item in segments]))
    return output_path, plan["overall"], student_text


def grade_homework(
    question_path: Path,
    student_path: Path,
    student_id: str,
    reference_path: Path | None,
    teacher_material_paths: list[Path] | None,
    protocol: str,
    api_key: str,
    base_url: str,
    model: str,
    output_dir: Path,
    use_model_file_inputs: bool = False,
) -> tuple[Path, str, dict[str, Any]]:
    ext = student_path.suffix.lower()
    if ext not in STUDENT_WORD_EXT:
        raise ValueError(f"待批改文件必须是Word格式(doc/docx)，当前为：{ext}")
    normalized_student_id = normalize_student_id(student_id)

    teacher_paths = [question_path] + (teacher_material_paths or [])
    use_file_mode = use_model_file_inputs and protocol != "Anthropic兼容"

    def _build_text_context(context_teacher_paths: list[Path], context_reference_path: Path | None) -> tuple[str, str]:
        teacher_text_parts = []
        for p in context_teacher_paths:
            teacher_text_parts.append(f"[教师材料: {p.name}]")
            teacher_text_parts.append(extract_text(p))
        question_txt = normalize_text("\n\n".join(teacher_text_parts))
        reference_txt = extract_text(context_reference_path) if context_reference_path else ""
        return question_txt, reference_txt

    if use_file_mode:
        file_teacher_paths, text_teacher_paths = split_file_api_supported(teacher_paths)
        file_reference_path = None
        text_reference_path = reference_path
        if reference_path and reference_path.suffix.lower() in FILE_API_SUPPORTED_EXT:
            file_reference_path = reference_path
            text_reference_path = None
        question_text, reference_text = _build_text_context(text_teacher_paths, text_reference_path)
        try:
            output_path, overall, student_text = annotate_word(
                student_path,
                normalized_student_id,
                output_dir,
                protocol,
                api_key,
                base_url,
                model,
                question_text,
                reference_text,
                question_file_paths=file_teacher_paths,
                reference_file_path=file_reference_path,
                use_model_file_inputs=True,
            )
            kp_event = generate_knowledge_point_results(
                protocol=protocol,
                api_key=api_key,
                base_url=base_url,
                model=model,
                student_id=normalized_student_id,
                question_text=question_text,
                student_text=student_text,
                reference_text=reference_text,
                question_file_paths=file_teacher_paths,
                reference_file_path=file_reference_path,
                use_model_file_inputs=True,
            )
        except Exception as exc:
            if not is_file_input_unsupported_error(exc):
                raise
            question_text, reference_text = _build_text_context(teacher_paths, reference_path)
            output_path, overall, student_text = annotate_word(
                student_path,
                normalized_student_id,
                output_dir,
                protocol,
                api_key,
                base_url,
                model,
                question_text,
                reference_text,
                question_file_paths=teacher_paths,
                reference_file_path=reference_path,
                use_model_file_inputs=False,
            )
            kp_event = generate_knowledge_point_results(
                protocol=protocol,
                api_key=api_key,
                base_url=base_url,
                model=model,
                student_id=normalized_student_id,
                question_text=question_text,
                student_text=student_text,
                reference_text=reference_text,
                question_file_paths=teacher_paths,
                reference_file_path=reference_path,
                use_model_file_inputs=False,
            )
    else:
        question_text, reference_text = _build_text_context(teacher_paths, reference_path)
        output_path, overall, student_text = annotate_word(
            student_path,
            normalized_student_id,
            output_dir,
            protocol,
            api_key,
            base_url,
            model,
            question_text,
            reference_text,
            question_file_paths=teacher_paths,
            reference_file_path=reference_path,
            use_model_file_inputs=False,
        )
        kp_event = generate_knowledge_point_results(
            protocol=protocol,
            api_key=api_key,
            base_url=base_url,
            model=model,
            student_id=normalized_student_id,
            question_text=question_text,
            student_text=student_text,
            reference_text=reference_text,
            question_file_paths=teacher_paths,
            reference_file_path=reference_path,
            use_model_file_inputs=False,
        )
    return output_path, overall, kp_event


def save_upload(uploaded_file, target_dir: Path) -> Path:
    target_dir.mkdir(parents=True, exist_ok=True)
    file_name = getattr(uploaded_file, "name", None) or getattr(uploaded_file, "filename", None)
    if not file_name:
        raise ValueError("上传文件缺少文件名。")
    target = target_dir / Path(file_name).name

    if hasattr(uploaded_file, "getbuffer"):
        target.write_bytes(uploaded_file.getbuffer())
        return target

    file_obj = getattr(uploaded_file, "file", None)
    if file_obj is not None and hasattr(file_obj, "read"):
        file_obj.seek(0)
        target.write_bytes(file_obj.read())
        return target

    raise ValueError("不支持的上传文件对象类型。")


def render_upload_summary(
    question_upload,
    teacher_material_uploads,
    student_upload,
    reference_upload,
    student_id: str,
) -> None:
    st.markdown("### 已上传文件")
    st.write(f"- 学生ID：`{student_id.strip() or '未填写'}`")
    if question_upload:
        st.write(f"- 题目：`{question_upload.name}`")
    else:
        st.write("- 题目：`未上传`")

    if teacher_material_uploads:
        for item in teacher_material_uploads:
            st.write(f"- 材料：`{item.name}`")
    else:
        st.write("- 材料：`无`")

    if student_upload:
        st.write(f"- 作业：`{student_upload.name}`")
    else:
        st.write("- 作业：`未上传`")

    if reference_upload:
        st.write(f"- 样例：`{reference_upload.name}`")
    else:
        st.write("- 样例：`无`")


def main() -> None:
    st.set_page_config(page_title="作业批注系统", page_icon="📝", layout="centered")
    st.title("作业批注前端页")
    st.caption("上传题目/材料/学生作业，系统处理后输出批注完成的 Word 文件")

    with st.sidebar:
        st.markdown("### 模型设置")
        protocol = st.selectbox("接口协议", ["OpenAI兼容", "Anthropic兼容"], index=0)
        api_key = st.text_input("ARK API Key", value=os.getenv("ARK_API_KEY", ""), type="password")
        model = st.text_input("模型名", value=os.getenv("ARK_MODEL", DEFAULT_MODEL))

        default_base_url = os.getenv("ARK_BASE_URL", DEFAULT_OPENAI_BASE_URL)
        if protocol == "Anthropic兼容":
            default_base_url = DEFAULT_ANTHROPIC_BASE_URL
        base_url = st.text_input("Base URL", value=default_base_url)
        use_model_file_inputs = st.checkbox(
            "直传文件给模型（实验）",
            value=True,
            help="启用后将题目/材料/样例文件作为附件直接发送给模型，不再在本地解析这些文件。",
        )
        member4_ingest_url = st.text_input(
            "成员4 Ingest URL",
            value=os.getenv("MEMBER4_INGEST_URL", DEFAULT_MEMBER4_INGEST_URL),
        )
        source_id = st.text_input("source_id", value=DEFAULT_SOURCE_ID, help="可使用 assignment-001 或任务ID")

        st.markdown("### 支持格式")
        st.write("- 题目/材料：`pdf/doc/docx/xls/xlsx`")
        st.write("- 学生作业：`doc/docx`")
        st.write("- 输出：批注后的 `docx` + 结构化知识点结果")
        st.caption("提示：Linux 环境处理 .doc/.xls 需要安装 LibreOffice（soffice）自动转换。")

    with st.form("grading_form"):
        st.markdown("### 0) 填写学生ID")
        student_id = st.text_input("学生ID（必填，仅用于标识和输出文件命名）")

        st.markdown("### 1) 上传题目")
        question_upload = st.file_uploader(
            "题目文件（必填）",
            type=["pdf", "doc", "docx", "xls", "xlsx"],
        )

        st.markdown("### 2) 上传老师材料（可选）")
        teacher_material_uploads = st.file_uploader(
            "补充材料（可多选）",
            type=["pdf", "doc", "docx", "xls", "xlsx"],
            accept_multiple_files=True,
        )

        st.markdown("### 3) 上传学生作业")
        student_upload = st.file_uploader("学生待批改 Word（必填）", type=["doc", "docx"])

        st.markdown("### 4) 上传参考样例（可选）")
        reference_upload = st.file_uploader(
            "老师批改样例",
            type=["pdf", "doc", "docx", "xls", "xlsx"],
        )

        submitted = st.form_submit_button("开始处理并生成批注Word", type="primary")

    render_upload_summary(question_upload, teacher_material_uploads, student_upload, reference_upload, student_id)

    if submitted:
        if not question_upload or not student_upload:
            st.error("请至少上传题目文件和学生待批改文件。")
            return
        if not student_id.strip():
            st.error("请填写学生ID。")
            return
        if not api_key.strip():
            st.error("请填写 API Key。")
            return

        question_ext = Path(question_upload.name).suffix.lower()
        student_ext = Path(student_upload.name).suffix.lower()
        if question_ext not in QUESTION_EXT:
            st.error("题目文件格式不支持，请使用 pdf/doc/docx/xls/xlsx。")
            return
        if student_ext not in STUDENT_WORD_EXT:
            st.error("学生待批改文件必须是 doc/docx。")
            return

        work_dir = Path("workspace_uploads") / now_suffix()

        with st.spinner("正在批改，请稍候..."):
            try:
                normalized_student_id = normalize_student_id(student_id)
                question_path = save_upload(question_upload, work_dir)
                teacher_material_paths = [save_upload(f, work_dir) for f in (teacher_material_uploads or [])]
                student_path = save_upload(student_upload, work_dir)
                reference_path = save_upload(reference_upload, work_dir) if reference_upload else None

                output_path, overall, kp_event = grade_homework(
                    question_path,
                    student_path,
                    normalized_student_id,
                    reference_path,
                    teacher_material_paths,
                    protocol,
                    api_key,
                    base_url,
                    model,
                    OUTPUT_DIR,
                    use_model_file_inputs=use_model_file_inputs,
                )
                kp_results = kp_event.get("payload", {}).get("results", [])
                ingest_url = member4_ingest_url.strip() or DEFAULT_MEMBER4_INGEST_URL
                push_result = post_member4_event(
                    ingest_url=ingest_url,
                    user_id=normalized_student_id,
                    source_id=source_id.strip() or DEFAULT_SOURCE_ID,
                    results=kp_results,
                )

                st.success("批改完成")
                st.write(f"总评：{overall}")
                st.code(f"输出文件：{output_path.resolve()}")
                st.markdown("### 知识点正确性结果")
                st.json(
                    {
                        "events": [
                            {
                                "user_id": normalized_student_id,
                                "class_id": CLASS_ID,
                                "course_id": COURSE_ID,
                                "event_type": EVENT_TYPE,
                                "payload": {
                                    "source_id": source_id.strip() or DEFAULT_SOURCE_ID,
                                    "source_type": SOURCE_TYPE,
                                    "results": kp_results,
                                },
                            }
                        ]
                    }
                )
                st.markdown("### 成员4推送结果")
                st.json(push_result)
                st.download_button(
                    "下载批改后文件",
                    data=output_path.read_bytes(),
                    file_name=output_path.name,
                    mime="application/octet-stream",
                )
            except Exception as exc:
                st.error(f"批改失败：{exc}")


if __name__ == "__main__":
    main()
